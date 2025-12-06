"""
Document Processor Service

Extracts text from various document formats (PDF, DOCX, TXT) and splits
into chunks suitable for embedding and retrieval.
"""

import json
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Extracts and chunks text from various document formats.

    Supports:
    - PDF (via PyMuPDF/fitz)
    - DOCX (via python-docx)
    - TXT/MD (plain text)
    """

    SUPPORTED_MIME_TYPES = {
        "application/pdf": "_process_pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "_process_docx",
        "text/plain": "_process_txt",
        "text/markdown": "_process_txt",
    }

    # File extension to MIME type mapping
    EXTENSION_MAP = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }

    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None,
    ):
        """
        Initialize the document processor.

        Args:
            chunk_size: Characters per chunk (default from settings)
            chunk_overlap: Overlap between chunks (default from settings)
        """
        self.chunk_size = chunk_size or settings.RAG_CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.RAG_CHUNK_OVERLAP

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def get_mime_type(self, file_path: str) -> Optional[str]:
        """Determine MIME type from file extension."""
        ext = Path(file_path).suffix.lower()
        return self.EXTENSION_MAP.get(ext)

    def is_supported(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if file type is supported."""
        if mime_type:
            return mime_type in self.SUPPORTED_MIME_TYPES
        if file_path:
            detected_mime = self.get_mime_type(file_path)
            return detected_mime in self.SUPPORTED_MIME_TYPES
        return False

    async def process_document(
        self,
        file_path: str,
        mime_type: str = None,
    ) -> List[Dict[str, Any]]:
        """
        Extract text from a document and split into chunks.

        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document (auto-detected if not provided)

        Returns:
            List of chunk dictionaries with content and metadata

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Auto-detect MIME type if not provided
        if not mime_type:
            mime_type = self.get_mime_type(file_path)

        if not mime_type or mime_type not in self.SUPPORTED_MIME_TYPES:
            raise ValueError(f"Unsupported file type: {mime_type or 'unknown'}")

        # Get the appropriate processor method
        processor_method = getattr(self, self.SUPPORTED_MIME_TYPES[mime_type])

        # Extract text
        logger.info(f"Processing document: {file_path} ({mime_type})")
        text, metadata = processor_method(file_path)

        if not text or not text.strip():
            logger.warning(f"No text extracted from document: {file_path}")
            return []

        # Split into chunks
        chunks = self._split_text(text, metadata, file_path)
        logger.info(f"Created {len(chunks)} chunks from document: {file_path}")

        return chunks

    def _process_pdf(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using PyMuPDF.

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF is required for PDF processing. Install with: pip install PyMuPDF")

        text_parts = []
        metadata = {
            "source": file_path,
            "file_type": "pdf",
            "page_count": 0,
        }

        try:
            doc = fitz.open(file_path)
            metadata["page_count"] = len(doc)
            metadata["title"] = doc.metadata.get("title", "")
            metadata["author"] = doc.metadata.get("author", "")

            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")

            doc.close()
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

        return "\n\n".join(text_parts), metadata

    def _process_docx(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX using python-docx.

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

        metadata = {
            "source": file_path,
            "file_type": "docx",
        }

        try:
            doc = Document(file_path)

            # Extract core properties
            if doc.core_properties:
                metadata["title"] = doc.core_properties.title or ""
                metadata["author"] = doc.core_properties.author or ""
                metadata["subject"] = doc.core_properties.subject or ""

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Include heading style as section markers
                    if para.style and para.style.name.startswith("Heading"):
                        text_parts.append(f"\n## {para.text}\n")
                    else:
                        text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

        return "\n".join(text_parts), metadata

    def _process_txt(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Read plain text files.

        Returns:
            Tuple of (text_content, metadata_dict)
        """
        metadata = {
            "source": file_path,
            "file_type": "txt" if file_path.endswith(".txt") else "md",
        }

        try:
            # Try different encodings
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
            text = None

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ValueError(f"Could not decode file: {file_path}")

        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise

        return text, metadata

    def _split_text(
        self,
        text: str,
        metadata: Dict[str, Any],
        file_path: str,
    ) -> List[Dict[str, Any]]:
        """
        Split text into chunks with metadata.

        Args:
            text: Full document text
            metadata: Document-level metadata
            file_path: Source file path

        Returns:
            List of chunk dictionaries
        """
        # Use LangChain's text splitter
        chunks = self.text_splitter.split_text(text)

        result = []
        for idx, chunk_content in enumerate(chunks):
            chunk_metadata = {
                **metadata,
                "chunk_index": idx,
                "total_chunks": len(chunks),
            }

            result.append({
                "content": chunk_content,
                "metadata": json.dumps(chunk_metadata),
                "chunk_index": idx,
            })

        return result

    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get basic information about a document without full processing.

        Args:
            file_path: Path to the document

        Returns:
            Dictionary with file information
        """
        path = Path(file_path)
        stat = path.stat()

        return {
            "file_name": path.name,
            "file_size": stat.st_size,
            "mime_type": self.get_mime_type(file_path),
            "is_supported": self.is_supported(file_path=file_path),
            "extension": path.suffix.lower(),
        }
